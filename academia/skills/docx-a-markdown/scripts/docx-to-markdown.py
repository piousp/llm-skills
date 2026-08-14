#!/usr/bin/env python3
"""
docx-to-markdown.py — Convierte un archivo DOCX a Markdown y extrae
opcionalmente los comentarios de Word.

Motores:
  pandoc        (default) conversión vía pandoc
  mammoth       conversión vía mammoth
  python-docx   conversión básica vía python-docx (fallback)

Uso:
  python3 docx-to-markdown.py <archivo.docx> [--output salida.md]
                               [--mode pandoc|mammoth|python-docx]
                               [--no-comments]
"""

import argparse
import json
import os
import subprocess
import sys
from pathlib import Path


def run(cmd: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, capture_output=True, text=True)


def convert_pandoc(docx: Path) -> str:
    result = run(["pandoc", str(docx), "-t", "markdown", "--wrap=preserve"])
    if result.returncode != 0:
        sys.exit(
            f"Error en pandoc (código {result.returncode}): "
            f"{result.stderr or result.stdout}"
        )
    return result.stdout


def convert_mammoth(docx: Path) -> str:
    result = run(["mammoth", str(docx), "--output-format=markdown"])
    if result.returncode != 0:
        sys.exit(
            f"Error en mammoth (código {result.returncode}): "
            f"{result.stderr or result.stdout}"
        )
    return result.stdout


def convert_python_docx(docx: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit("python-docx no está instalado (pip install python-docx)")
    doc = Document(str(docx))
    out: list[str] = []
    for p in doc.paragraphs:
        out.append(p.text + "\n\n")
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            out.append("| " + " | ".join(cells) + " |\n")
        out.append("\n")
    return "".join(out)


def extract_comments(docx: Path, script: Path) -> str:
    """Ejecuta extract-comments.py y devuelve la sección Markdown, o ''."""
    result = run(["python3", str(script), str(docx)])
    if result.returncode != 0 or not result.stdout.strip():
        return ""
    try:
        comments = json.loads(result.stdout)
    except json.JSONDecodeError:
        return ""
    if not isinstance(comments, list) or not comments:
        return ""

    lines = ["", "---", "", "## Comentarios del documento", ""]
    for c in comments:
        lines.append(f"### Comentario #{c.get('id')}")
        lines.append("")
        lines.append(f"- **Autor:** {c.get('author')}")
        if c.get("date"):
            lines.append(f"- **Fecha:** {c.get('date')}")
        if c.get("annotated_text"):
            lines.append(f'- **Sobre el texto:** "{c.get("annotated_text")}"')
        lines.append("")
        lines.append(c.get("text", ""))
        lines.append("")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convierte DOCX a Markdown.")
    parser.add_argument("path", help="Ruta al archivo .docx")
    parser.add_argument(
        "--output",
        help="Ruta del .md de salida (si se omite, imprime por stdout)",
    )
    parser.add_argument(
        "--mode",
        choices=["pandoc", "mammoth", "python-docx"],
        default="pandoc",
        help="Motor de conversión (default: pandoc)",
    )
    parser.add_argument(
        "--no-comments",
        action="store_true",
        help="No extraer comentarios de Word (default: incluirlos)",
    )
    args = parser.parse_args()

    docx = Path(args.path)
    if not docx.exists():
        sys.exit(f"No existe el archivo: {docx}")

    if args.mode == "pandoc":
        md_content = convert_pandoc(docx)
    elif args.mode == "mammoth":
        md_content = convert_mammoth(docx)
    else:
        md_content = convert_python_docx(docx)

    comments_section = ""
    if not args.no_comments:
        script = Path(__file__).resolve().parent / "extract-comments.py"
        comments_section = extract_comments(docx, script)

    final_content = md_content.rstrip() + comments_section + "\n"

    if args.output:
        output_path = Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(final_content, encoding="utf-8")
        print(f"Documento convertido: {output_path}")
    else:
        sys.stdout.write(final_content)


if __name__ == "__main__":
    try:
        main()
    except BrokenPipeError:
        # Salir silenciosamente si stdout se cierra antes de tiempo (p. ej. `| head`)
        devnull = os.open(os.devnull, os.O_WRONLY)
        os.dup2(devnull, sys.stdout.fileno())
        sys.exit(1)
